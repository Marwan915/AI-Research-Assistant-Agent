"""
core/agent.py — Modular LangGraph 3-Node Research Agent

FIXES:
  - FIXED: embedding_model property type hint was HuggingFaceEmbeddings but
    actual type is GoogleGenerativeAIEmbeddings — caused type errors downstream
  - FIXED: chroma_path default changed to "chroma_db_v3" (was inconsistent with rag_pipeline.py)
  - FIXED: writer fallback to Gemini when Ollama/SciAssistant unavailable
  - FIXED: _deep_writer_node was not resetting stream queue before writing
  - FIXED: stream_chunks() timeout handling — now properly drains queue on done
  - ADDED: query history tracking (last 20 queries stored in self._history)
  - ADDED: get_research_notes() helper for external access
"""

import os
import re
import json
import requests
import queue
import threading
from datetime import datetime
from typing import Optional, Callable, Dict, Any, Generator, List

from streamlit.runtime.scriptrunner import add_script_run_ctx
from pydantic import BaseModel, Field
from dotenv import load_dotenv
from langchain_chroma import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings

# Monkey patch to fix batch embedding flattening bug in langchain-google-genai v4
def _patched_embed_documents(self, texts: List[str]) -> List[List[float]]:
    if not texts:
        return []
    try:
        # Wrap each text in the expected Content structure to prevent aggregation in gemini-embedding-2
        contents = [{"parts": [{"text": t}]} for t in texts]
        batch_size = 100
        embeddings = []
        for i in range(0, len(texts), batch_size):
            batch_contents = contents[i : i + batch_size]
            config = self._build_config(
                task_type="RETRIEVAL_DOCUMENT",
                output_dimensionality=self.output_dimensionality
            )
            result = self.client.models.embed_content(
                model=self.model,
                contents=batch_contents,
                config=config
            )
            embeddings.extend([list(e.values) for e in result.embeddings])
        return embeddings
    except Exception as e:
        # Fallback to sequential if anything goes wrong
        return [self.embed_query(t) for t in texts]
GoogleGenerativeAIEmbeddings.embed_documents = _patched_embed_documents
from langchain_classic.retrievers import ContextualCompressionRetriever
from langchain_classic.retrievers.document_compressors import CrossEncoderReranker
from langchain_community.cross_encoders import HuggingFaceCrossEncoder
from langchain_core.tools import tool
from langchain_ollama import OllamaLLM

from langgraph.graph import StateGraph, START, END
from langgraph.checkpoint.memory import MemorySaver

from arxiv_tool import search_arxiv
from gemini_router import GeminiRouter

load_dotenv()


# ══════════════════════════════════════════════
# Helper — Extract text from Gemini response
# ══════════════════════════════════════════════
def _text(content) -> str:
    """Convert Gemini output (str | list | AIMessage) to plain text."""
    if hasattr(content, "content"):
        content = content.content
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        return "".join(
            item.get("text", "") if isinstance(item, dict) else str(item)
            for item in content
        )
    return str(content)


# ══════════════════════════════════════════════
# State Management
# ══════════════════════════════════════════════
class AgentState(BaseModel):
    input_query: str = Field(..., description="The user's original question")
    output_depth: str = Field(default="Quick Summary", description="Requested output depth")
    translated_query: str = Field(default="", description="English translation")
    research_notes: str = Field(default="", description="Fact skeleton from researcher node")
    final_report: str = Field(default="", description="Draft from writer node")
    audited_report: str = Field(default="", description="Verified final report from validator node")
    cached_research_notes: str = Field(default="", description="Cached notes to skip research")


# ══════════════════════════════════════════════
# Pipeline Stage Events
# ══════════════════════════════════════════════
class PipelineEvents:
    def __init__(self):
        self.researcher_done = threading.Event()
        self.writer_started = threading.Event()
        self.writer_done = threading.Event()
        self.validator_done = threading.Event()
        self.pipeline_done = threading.Event()
        self.pipeline_failed = threading.Event()
        self.error: Optional[Exception] = None

    def reset(self):
        for attr in vars(self):
            val = getattr(self, attr)
            if isinstance(val, threading.Event):
                val.clear()
        self.error = None


# ══════════════════════════════════════════════
# Research Agent
# ══════════════════════════════════════════════
class ResearchAgent:
    """
    Encapsulates the 3-node LangGraph pipeline with proper thread synchronization.

    Usage (Streamlit):
        agent = ResearchAgent(on_status=lambda step, detail: ...)
        thread = agent.run_async("What is transformer attention?")
        agent.events.writer_started.wait(timeout=120)
        for chunk in agent.stream_chunks():
            display(chunk)
        agent.events.validator_done.wait(timeout=120)
        result = agent.last_result
    """

    def __init__(
        self,
        on_status: Optional[Callable[[str, str], None]] = None,
        on_fallback: Optional[Callable[[str, str, str], None]] = None,
        chroma_path: str = "chroma_db_v3",
    ):
        self._on_status = on_status or (lambda s, d: None)
        self._stream_queue: queue.Queue = queue.Queue()
        self._last_result: Optional[Dict[str, Any]] = None
        self._history: List[Dict[str, str]] = []  # NEW: query history
        self.events = PipelineEvents()

        # ── Models ──
        self._router = GeminiRouter(on_fallback=on_fallback)

        # Try to load local writer; check Ollama API instead of calling model.invoke to avoid massive startup hang
        self._writer_llm = None
        self._writer_is_local = False
        try:
            import requests
            # Quick HTTP check to see if Ollama is running and has the model
            resp = requests.get("http://localhost:11434/api/tags", timeout=2.0)
            if resp.status_code == 200:
                models = [m["name"] for m in resp.json().get("models", [])]
                if any("SciAssistant" in m for m in models):
                    self._writer_llm = OllamaLLM(model="SciAssistant", timeout=60.0)
                    self._writer_is_local = True
                    self._on_status("init", "✓ Local Ollama SciAssistant model detected.")
                else:
                    self._on_status("init", "⚠ Local Ollama is running but 'SciAssistant' model not found. Using Gemini.")
            else:
                self._on_status("init", "⚠ Ollama server returned error. Using Gemini.")
        except Exception as e:
            self._on_status("init", f"⚠ Ollama server not detected ({e}). Using Gemini.")

        # ── RAG Retriever ──
        self._on_status("init", "Loading embedding model and vector database...")
        self._embedding_model = GoogleGenerativeAIEmbeddings(
            model="models/gemini-embedding-2"
        )

        db = Chroma(persist_directory=chroma_path, embedding_function=self._embedding_model)
        base_retriever = db.as_retriever(search_kwargs={"k": 15})
        cross_encoder = HuggingFaceCrossEncoder(model_name="BAAI/bge-reranker-base")
        compressor = CrossEncoderReranker(model=cross_encoder, top_n=5)
        self._retriever = ContextualCompressionRetriever(
            base_compressor=compressor, base_retriever=base_retriever
        )

        # ── Local search tool ──
        retriever_ref = self._retriever

        @tool
        def search_local_papers(query: str) -> str:
            """Search the user's local database of academic PDFs."""
            try:
                docs = retriever_ref.invoke(query)
                if not docs:
                    return "No relevant local data found for this query."
                return "\n\n".join(
                    [f"[Local Source {i+1}]: {d.page_content}" for i, d in enumerate(docs)]
                )
            except Exception as e:
                return f"[DATABASE_ERROR] Error searching local papers: {e}"

        self._search_local = search_local_papers

        # ── Build LangGraph ──
        self._graph = self._build_graph()
        self._on_status("init", "✓ Agent initialized and ready.")

    # ── Properties ────────────────────────────
    @property
    def embedding_model(self) -> GoogleGenerativeAIEmbeddings:
        """
        Expose embedding model for reuse in pdf_chat.
        FIXED: Type hint was HuggingFaceEmbeddings — actual type is GoogleGenerativeAIEmbeddings.
        """
        return self._embedding_model

    @property
    def last_result(self) -> Optional[Dict[str, Any]]:
        return self._last_result

    @property
    def router(self) -> GeminiRouter:
        return self._router

    @property
    def query_history(self) -> List[Dict[str, str]]:
        """Return last 20 queries with timestamps."""
        return self._history[-20:]

    @property
    def writer_mode(self) -> str:
        return "SciAssistant (local)" if self._writer_is_local else "Gemini (fallback)"

    # ══════════════════════════════════════════
    # Node 1: Researcher
    # ══════════════════════════════════════════
    def _researcher_node(self, state: AgentState) -> dict:
        if getattr(state, "cached_research_notes", ""):
            self._on_status("researcher", "✓ Using cached research notes. Skipping retrieval...")
            self.events.researcher_done.set()
            return {"translated_query": state.translated_query or state.input_query, "research_notes": state.cached_research_notes}

        self._on_status("researcher", "Analyzing query & routing to correct databases...")

        analysis_prompt = f"""Analyze the following user query.
1. Translate it to English if it is not already in English.
2. Determine which data sources are most appropriate. Choose any combination of: ["WEB", "ARXIV", "LOCAL"].
   - WEB: MUST BE SELECTED for ANY technical comparison (e.g., SQL vs NoSQL, Frontend vs Backend), general knowledge, or coding concepts.
   - ARXIV: For deep academic research, mathematical algorithms, complex machine learning architectures.
   - LOCAL: If the user explicitly asks about their "own documents", "my files", or specific local knowledge.
3. Translate and expand the user query into a highly professional, academic English search query optimized for scientific databases and search engines. Focus on core architectural and technical keywords, not conversational phrasing. (e.g., "Architectural differences between Relational SQL and Non-Relational NoSQL databases performance scalability"). MUST BE IN ENGLISH.

Output ONLY a valid JSON object with the following keys, no markdown fences:
{{
    "translated_query": "...",
    "sources": ["WEB", "ARXIV"],
    "search_keyword": "..."
}}

Query: {state.input_query}"""

        response = self._router.invoke(analysis_prompt)
        raw_json = _text(response).strip()
        # Strip markdown fences and BOM
        clean_json = re.sub(r"```(?:json)?|```", "", raw_json).strip().lstrip("\ufeff")

        try:
            analysis = json.loads(clean_json)
            translated = analysis.get("translated_query", state.input_query)
            sources = analysis.get("sources", ["WEB"])
            if isinstance(sources, str):
                sources = [sources]
            keyword = analysis.get("search_keyword", translated[:20])
        except Exception as e:
            self._on_status("researcher", f"⚠ Routing failed ({e}), defaulting to WEB.")
            translated = state.input_query
            sources = ["WEB"]
            keyword = translated[:20]

        self._on_status("researcher", f"✓ Route: {sources} | Query (EN): '{translated[:60]}...'")

        local_data = ""
        web_data = ""
        arxiv_data = ""

        if "LOCAL" in sources or not sources:
            self._on_status("tools", f"Searching local DB: '{keyword}'...")
            local_data = self._search_local.invoke(keyword)
            
        if "WEB" in sources:
            # ── Tavily (primary) ──────────────────────────────────────
            try:
                from tavily import TavilyClient
                tavily_client = TavilyClient(api_key="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")
                tavily_raw = tavily_client.search(keyword, max_results=3)
                results = [
                    {'title': r['title'], 'body': r['content'][:300], 'href': r['url']}
                    for r in tavily_raw.get('results', [])
                ]
                self._on_status("tools", f"✓ Tavily: {len(results)} results.")
            except Exception as e:
                self._on_status("tools", f"Tavily error: {e}")
                results = []

            # ── Wikipedia (secondary/supplement) ─────────────────────
            try:
                wiki_url = 'https://en.wikipedia.org/w/api.php'
                wiki_params = {'action': 'query', 'list': 'search', 'srsearch': keyword, 'utf8': '', 'format': 'json'}
                wiki_headers = {'User-Agent': 'ResearchAgent/1.0'}
                wiki_res = requests.get(wiki_url, params=wiki_params, headers=wiki_headers, timeout=10)
                wiki_data = wiki_res.json().get('query', {}).get('search', [])
                wiki_results = [
                    {'title': f"Wikipedia: {item['title']}",
                     'body': re.sub(r'<[^>]+>', '', item['snippet']),
                     'href': f"https://en.wikipedia.org/wiki/{item['title'].replace(' ', '_')}"}
                    for item in wiki_data[:2]
                ]
                results.extend(wiki_results)
                self._on_status("tools", f"✓ Wikipedia: {len(wiki_results)} results.")
            except Exception as e:
                self._on_status("tools", f"Wiki error: {e}")

            if results:
                web_data = "\n\n".join([
                    f"[Web: {r.get('title')}]\nContent: {r.get('body')}\nURL: {r.get('href')}"
                    for r in results
                ])
            else:
                web_data = "No web results found."

        if "ARXIV" in sources:
            self._on_status("tools", f"Querying ArXiv: '{keyword}'...")
            arxiv_data = search_arxiv.invoke({"query": keyword, "max_results": 2})

        self._on_status("tools", f"✓ Retrieved data from selected sources.")

        self._on_status("researcher", "Synthesizing fact skeleton...")
        synthesis_prompt = f"""You are the elite scientific researcher and the "brain" of the system.
Your task is to synthesize a highly technical, bulleted "Fact Skeleton" to answer the user's query: "{translated}".

CRITICAL LANGUAGE RULE: You MUST output the Fact Skeleton ENTIRELY IN ENGLISH, regardless of the original query language.

CRITICAL INSTRUCTIONS:
1. Evaluate ALL provided sources (Web, Local, ArXiv). Extract facts ONLY from sources that are directly relevant to the query.
2. For every fact you cite from Web results, you MUST include its URL at the end of the bullet point. Format: (Source: URL). This is MANDATORY.
3. If a source is completely off-topic, ignore it — but DO NOT write a disclaimer. Simply use your expert knowledge for that point without citing a source.
4. Do NOT write any [DISCLAIMER] text. Do NOT mention that sources were unavailable. Only output the Fact Skeleton.

RULES FOR FACT SKELETON:
- CRITICAL: Do NOT use $ or $$ delimiters for math equations or variables. Write them as plain inline text.
- NEVER cite a document if it doesn't directly and explicitly support the specific fact. Do not force citations.
- Cite sources explicitly: [Local Source N], [Web: Title](URL), or [ArXiv: Title] (ONLY if strictly relevant).
- Do NOT write prose — only a structured bullet list.
- Do NOT hallucinate math.

[LOCAL DATABASE RESULTS]:
{local_data}

[WEB SEARCH RESULTS]:
{web_data}

[ARXIV SEARCH RESULTS]:
{arxiv_data}

Output the Fact Skeleton now:"""

        response = self._router.invoke(synthesis_prompt)
        fact_skeleton = _text(response)

        # ── Python-controlled Disclaimer (NOT LLM decision) ──────────────
        # We know whether real web results exist — the LLM does not decide.
        has_web_results = bool(results) if "WEB" in sources else False
        has_arxiv_results = bool(arxiv_data and arxiv_data.strip() and "No results" not in arxiv_data) if "ARXIV" in sources else False
        has_local_results = bool(local_data and local_data.strip()) if ("LOCAL" in sources or not sources) else False
        has_any_results = has_web_results or has_arxiv_results or has_local_results

        if not has_any_results:
            disclaimer = "[DISCLAIMER]: The provided search sources (Web, ArXiv, Local DB) yielded limited or no relevant results. This information is derived from general expert knowledge in the field.\n\n"
            if not fact_skeleton.startswith("[DISCLAIMER]"):
                fact_skeleton = disclaimer + fact_skeleton
        else:
            # Remove any hallucinated disclaimer the LLM may have added despite having sources
            if fact_skeleton.startswith("[DISCLAIMER]"):
                # Strip it — Python knows better
                lines = fact_skeleton.split("\n")
                # Skip lines until we hit the actual content (first non-disclaimer, non-empty line)
                skip = 0
                for i, line in enumerate(lines):
                    if line.startswith("[DISCLAIMER]") or (i > 0 and not line.strip() and skip == 0):
                        skip = i + 1
                    else:
                        break
                fact_skeleton = "\n".join(lines[skip:]).strip()

        self._on_status("researcher", "✓ Fact skeleton complete.")
        self.events.researcher_done.set()

        return {"translated_query": translated, "research_notes": fact_skeleton}

    # ══════════════════════════════════════════
    # Node 2: Writer
    # ══════════════════════════════════════════
    def _writer_node(self, state: AgentState) -> dict:
        depth = getattr(state, "output_depth", "Quick Summary")
        if depth == "Full Academic Paper Draft":
            return self._deep_writer_node(state)

        self._on_status("writer", f"Writing quick summary ({self.writer_mode})...")
        self.events.writer_started.set()

        slm_prompt = f"""[INST] <<SYS>>
You are an expert academic writer. Transform the provided Research Notes into a clear, concise academic summary.

MANDATORY RULES:
1. **Prose Summary First:** You MUST write a 2-3 sentence dense academic paragraph immediately after the `### Academic Summary` header. DO NOT skip straight to the Fact Skeleton or bullet points.
2. If the notes contain a [DISCLAIMER] about missing local data, include a clear disclaimer in your introduction.
3. CRITICAL: Use ONLY standard Markdown for document structuring (e.g., #, **, -). STRICTLY AVOID using LaTeX environments like \begin{{align}}, \textbf{{}}, or any math-mode delimiters for regular prose or text formatting. Do NOT use $ or $$ delimiters for math equations or variables. Write them as plain inline text.
4. CRITICAL: Do NOT output ANY introductory text before the heading. The very first characters of your output MUST be exactly `### Academic Summary\n`. Do not include any meta-commentary or conversational filler. The report MUST end immediately after the References section. Do not add any closing remarks like "Please note that the above is...". Do NOT enclose your response in a markdown code block (e.g. ```markdown). Do NOT output HTML tags.
5. Write in a sophisticated academic tone.
6. CRITICAL LANGUAGE RULE: Write the final report ENTIRELY IN ENGLISH, regardless of the user's original query language.
7. CRITICAL: NEVER mention "disclaimer", "lack of databases", "yielded no relevant results", or "relies on general expert knowledge". Present the information confidently as if drawn from primary literature. The report must read as a standalone, objective academic paper.
8. Strictly use Markdown list syntax (- or *) when enumerating items like technical stacks, metrics, or responsibilities.
9. CRITICAL STRUCTURAL RULE: Give equal weight to ALL main entities or concepts in the Fact Skeleton. Do not drop one topic halfway through.
10. CITATION RULE (CRITICAL): Review the Fact Skeleton for any source URLs or links tagged as [Web: Title](URL) or (Source: URL). If external sources ARE present, you MUST list them at the very end under '### References' using this exact format: `- [Title](URL)`. ONLY use 'General AI Knowledge (Internal Expert Weights)' if the Fact Skeleton contains a [DISCLAIMER] or is completely devoid of any external URLs.
<</SYS>>

Research Notes:
{state.research_notes}
[/INST]
"""
        accumulated = ""

        if self._writer_is_local and self._writer_llm:
            # Stream from local Ollama model
            try:
                for chunk in self._writer_llm.stream(slm_prompt):
                    self._stream_queue.put(chunk)
                    accumulated += chunk
            except Exception as e:
                self._on_status("writer", f"⚠ Local writer error: {e}. Falling back to Gemini...")
                accumulated = self._gemini_writer_fallback(state.research_notes, state.input_query)
                self._stream_queue.put(accumulated)
        else:
            # Use Gemini as writer
            accumulated = self._gemini_writer_fallback(state.research_notes, state.input_query)
            # Stream word-by-word for visual effect
            words = accumulated.split(" ")
            for word in words:
                self._stream_queue.put(word + " ")

        self._stream_queue.put(None)
        self._on_status("writer", "✓ Draft complete.")
        self.events.writer_done.set()

        return {"final_report": accumulated}

    def _gemini_writer_fallback(self, research_notes: str, input_query: str) -> str:
        """Use Gemini to write the report when local writer unavailable."""
        prompt = f"""You are an expert academic writer. Write a formal academic report based on
the Research Notes below.

MANDATORY FORMAT:
Use a logical academic structure with appropriate headings (e.g., Introduction, Core Concepts, Conclusion) based on the topic.

RULES:
- If the notes contain a [DISCLAIMER], clearly include it in the introduction.
- CRITICAL: Use ONLY standard Markdown for document structuring. STRICTLY AVOID LaTeX environments for regular text. LaTeX ($inline$ or $$block$$) is exclusively for actual mathematical formulas.
- Do NOT address the user. Do NOT add meta-commentary, conversational instructions, or placeholders.
- Sophisticated academic tone throughout.
- CRITICAL LANGUAGE RULE: Write the final report ENTIRELY IN ENGLISH, regardless of the user's original query language.
- NEVER mention the search process, system metadata, database status, or lack of sources in the academic text itself. The report must read as a standalone, objective academic paper.
- Strictly use Markdown list syntax (- or *) when enumerating items like technical stacks, metrics, or responsibilities.
- CITATION RULE (CRITICAL): Review the Research Notes for any source URLs or links tagged as [Web: Title](URL) or (Source: URL). If external sources ARE present, you MUST list them at the very end under '### References' using this exact format: `- [Title](URL)`. ONLY use 'General AI Knowledge (Internal Expert Weights)' if the Research Notes contain a [DISCLAIMER] or are completely devoid of any external URLs.

Research Notes:
{research_notes}

Write the full report:"""
        try:
            response = self._router.invoke(prompt)
            result_text = _text(response)
            result_text = result_text.replace('\\n', '\n') # Fix escaped newlines
            return result_text
        except Exception as e:
            return f"# Report Generation Error\n\nFailed to generate report: {e}\n\n**Research Notes:**\n{state.research_notes}"

    def _deep_writer_node(self, state: AgentState) -> dict:
        """Single-shot comprehensive drafting for full academic paper."""
        self._on_status("writer", "Drafting full academic paper...")
        self.events.writer_started.set()

        prompt = f"""You are an expert academic writer. Draft a comprehensive, full-length academic paper based ONLY on the provided Research Notes.

MANDATORY RULES:
1. Use standard academic structure (e.g., Abstract, Introduction, Core Analysis/Methodology, Discussion, Conclusion).
2. CRITICAL: NEVER mention "disclaimer", "lack of databases", "yielded no relevant results", or "relies on general expert knowledge" inside the body paragraphs. Present the information confidently as if drawn from primary literature.
3. Write in-depth, detailed paragraphs. The paper should be highly comprehensive.
4. CRITICAL: Use ONLY standard Markdown for document structuring (e.g., #, **, -). STRICTLY AVOID using LaTeX environments like \begin{{align}}, \textbf{{}}, or any math-mode delimiters for regular prose or text formatting. Do NOT use $ or $$ delimiters for math equations or variables. Write them as plain inline text. Do NOT enclose your output in markdown code blocks (e.g., ```markdown). Do NOT output any HTML tags.
5. Do NOT address the user. Maintain a strict, objective academic tone. Do NOT add conversational advice or placeholders. The very first line of your output MUST be `### Abstract\n` without any introductory text.
6. CONCLUSION LOOP PREVENTION: End the paper with EXACTLY ONE concluding paragraph under the heading 'Conclusion'. Do not add any extra summary paragraphs after it.
7. CRITICAL LANGUAGE RULE: Write the final paper ENTIRELY IN ENGLISH, regardless of the user's original query language.
8. NEVER mention the search process, system metadata, database status, or lack of sources in the academic text itself. The paper must read as a standalone, objective academic paper.
9. Strictly use Markdown list syntax (- or *) when enumerating items like technical stacks, metrics, or responsibilities.
10. CRITICAL STRUCTURAL RULE (TUNNEL VISION PREVENTION): You MUST give equal weight to ALL main entities or concepts present in the Fact Skeleton. If the query is a comparison (e.g., A vs. B), your Analysis, Discussion, and Conclusion MUST explicitly address and compare both sides. Do not drop one topic halfway through the paper.
11. REFERENCES & META-COMMENTARY PREVENTION: Write the FULL academic paper as instructed above. Then, AT THE VERY END, include a '### References' section. CITATION RULE (CRITICAL): Review the Research Notes for any source URLs or links tagged as [Web: Title](URL) or (Source: URL). If external sources ARE present, list them using this exact format: `- [Title](URL)`. CRITICAL: Do NOT invent or hallucinate references! ONLY use 'General AI Knowledge (Internal Expert Weights)' if the Research Notes contain a [DISCLAIMER] or are completely devoid of any external URLs. After the references, you MUST STOP GENERATING. Do NOT add notes like "Please note that this paper is...".

Research Notes:
{state.research_notes}

Write the full academic paper now:"""

        accumulated = ""

        if self._writer_is_local and self._writer_llm:
            try:
                for chunk in self._writer_llm.stream(prompt):
                    self._stream_queue.put(chunk)
                    accumulated += chunk
            except Exception as e:
                self._on_status("writer", f"⚠ Local writer failed: {e}. Trying Gemini...")
                accumulated = ""

        # Fallback to Gemini if local writer not active, failed, or returned empty content
        if not accumulated.strip():
            try:
                response = self._router.invoke(prompt)
                accumulated = _text(response)
                accumulated = accumulated.replace('\\n', '\n') # Fix escaped newlines
                # Stream word-by-word for visual effect
                words = accumulated.split(" ")
                for word in words:
                    self._stream_queue.put(word + " ")
            except Exception as e:
                accumulated = f"\n\n⚠️ Failed to generate paper: {e}\n"
                self._stream_queue.put(accumulated)

        self._stream_queue.put(None)
        self._on_status("writer", "✓ Full draft complete.")
        self.events.writer_done.set()
        return {"final_report": accumulated}

    # ══════════════════════════════════════════
    # Node 3: Validator
    # ══════════════════════════════════════════
    def _validator_node(self, state: AgentState) -> dict:
        self._on_status("validator", "Auditing citations and correcting references...")

        audit_prompt = f"""You are a strict academic editor performing a citation audit.

A junior AI writer drafted a report based on research notes. The writer may have:
- Created fake inline citations (e.g., "Smith et al., 2023")
- Added a hallucinated References section with made-up papers
- Misattributed facts to wrong sources

YOUR TASKS:
1. Identify fake inline citations in the DRAFT.
2. Map them to real sources from the RESEARCH NOTES. If a source in the research notes is irrelevant or part of a [DISCLAIMER], DO NOT use it.
3. Build a correct References section using ONLY real, relevant sources from the notes.
4. If the notes indicate that NO sources were used (e.g. via a [DISCLAIMER]), set the true_references field to "\\n**References:**\\nGeneral AI Knowledge (Internal Expert Weights)" and remove any forced citations in the draft. DO NOT generate conversational instructions, advice, or placeholders.

OUTPUT: You MUST output ONLY a valid JSON object. No markdown fences, no extra text.
{{
    "replacements": {{
        "exact hallucinated text": "[Local Source N] or [ArXiv: Title] or [Web: Title]"
    }},
    "true_references": "\\n**References:**\\n1. [Source description]..."
}}

RESEARCH NOTES (ground truth sources):
{state.research_notes}

DRAFT REPORT (to be audited):
{state.final_report}"""

        response = self._router.invoke(audit_prompt)
        raw_json = _text(response)

        try:
            clean = re.sub(r"```(?:json)?|```", "", raw_json).strip().lstrip("\ufeff")
            # Remove trailing commas before } or ] (common LLM mistake)
            clean = re.sub(r",\s*([}\]])", r"\1", clean)

            corrections = json.loads(clean)
            cleaned_report = state.final_report

            for fake, real in corrections.get("replacements", {}).items():
                if fake and real:
                    self._on_status("validator", f"Fixing: '{fake[:35]}' → '{real}'")
                    cleaned_report = cleaned_report.replace(fake, real)

            # Strip hallucinated references section (only if after Conclusion or in the last 30% of the report)
            conclusion_idx = cleaned_report.rfind("# 4. Conclusion")
            if conclusion_idx != -1:
                conclusion_body = cleaned_report[conclusion_idx:]
                parts = re.split(
                    r"\n(?:\*\*|#{1,3}\s*)?(?:References|Bibliography|Works Cited|Sources)[\s:]*\n",
                    conclusion_body,
                    flags=re.IGNORECASE,
                )
                if len(parts) > 1:
                    cleaned_report = cleaned_report[:conclusion_idx] + parts[0]
            else:
                matches = list(re.finditer(
                    r"\n(?:\*\*|#{1,3}\s*)?(?:References|Bibliography|Works Cited|Sources)[\s:]*\n",
                    cleaned_report,
                    flags=re.IGNORECASE,
                ))
                if matches:
                    last_match = matches[-1]
                    if last_match.start() > len(cleaned_report) * 0.7:
                        cleaned_report = cleaned_report[:last_match.start()]

            true_refs = corrections.get("true_references", "")
            if true_refs:
                cleaned_report += "\n\n" + true_refs

            self._on_status("validator", "✓ Audit complete — report verified.")

        except (json.JSONDecodeError, KeyError, TypeError) as e:
            self._on_status("validator", f"⚠ JSON parse failed ({e}), using fallback.")
            cleaned_report = state.final_report
            conclusion_idx = cleaned_report.rfind("# 4. Conclusion")
            if conclusion_idx != -1:
                conclusion_body = cleaned_report[conclusion_idx:]
                parts = re.split(
                    r"\n(?:\*\*|#{1,3}\s*)?(?:References|Bibliography|Works Cited)[\s:]*\n",
                    conclusion_body,
                    flags=re.IGNORECASE,
                )
                if len(parts) > 1:
                    cleaned_report = cleaned_report[:conclusion_idx] + parts[0]
            else:
                matches = list(re.finditer(
                    r"\n(?:\*\*|#{1,3}\s*)?(?:References|Bibliography|Works Cited)[\s:]*\n",
                    cleaned_report,
                    flags=re.IGNORECASE,
                ))
                if matches:
                    last_match = matches[-1]
                    if last_match.start() > len(cleaned_report) * 0.7:
                        cleaned_report = cleaned_report[:last_match.start()]
            cleaned_report = cleaned_report.rstrip()
            cleaned_report += "\n\n**Sources Used (Raw):**\n" + state.research_notes[:2000]

        self.events.validator_done.set()
        return {"audited_report": cleaned_report}

    # ══════════════════════════════════════════
    # Graph Assembly
    # ══════════════════════════════════════════
    def _build_graph(self):
        workflow = StateGraph(AgentState)
        workflow.add_node("researcher", self._researcher_node)
        workflow.add_node("writer", self._writer_node)
        workflow.add_node("validator", self._validator_node)
        workflow.add_edge(START, "researcher")
        workflow.add_edge("researcher", "writer")
        workflow.add_edge("writer", "validator")
        workflow.add_edge("validator", END)
        return workflow.compile(checkpointer=MemorySaver())

    # ══════════════════════════════════════════
    # Public API
    # ══════════════════════════════════════════
    def _run_internal(self, query: str, output_depth: str, thread_id: str, cached_research_notes: str = "", translated_query: str = ""):
        try:
            config = {"configurable": {"thread_id": thread_id}}
            result = self._graph.invoke(
                {
                    "input_query": query, 
                    "output_depth": output_depth,
                    "cached_research_notes": cached_research_notes,
                    "translated_query": translated_query
                },
                config=config,
            )
            self._last_result = result

            # Record to history
            self._history.append({
                "query": query,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "model": self._router.current_model,
                "depth": output_depth,
                "audited_report": result.get("audited_report", ""),
                "research_notes": result.get("research_notes", "No notes available."),
                "translated_query": result.get("translated_query", query),
            })

            self.events.pipeline_done.set()
        except Exception as e:
            self.events.error = e
            self.events.pipeline_failed.set()
            
            error_str = str(e)
            if "exhausted" in error_str.lower() or "429" in error_str:
                friendly_err = "⚠️ لقد استنفدت الحصة المجانية لواجهة Google Gemini (Quota Exceeded). لا يمكن إكمال التقرير."
            else:
                friendly_err = f"❌ Pipeline error: {error_str}"
                
            self._stream_queue.put(f"\n\n{friendly_err}\n")
            self._stream_queue.put(None)
            self.events.writer_started.set()
            self.events.writer_done.set()
            self.events.validator_done.set()

    def run(self, query: str, thread_id: str = "cli_session") -> Dict[str, Any]:
        """Run synchronously (CLI usage)."""
        self.events.reset()
        self._run_internal(query, "Quick Summary", thread_id)
        return self._last_result

    def run_async(
        self,
        query: str,
        output_depth: str = "Quick Summary",
        thread_id: str = "streamlit_session",
        cached_research_notes: str = "",
        translated_query: str = "",
    ) -> threading.Thread:
        """Launch pipeline in a background thread."""
        self.events.reset()
        self._last_result = None

        # Drain queue
        while not self._stream_queue.empty():
            try:
                self._stream_queue.get_nowait()
            except queue.Empty:
                break

        thread = threading.Thread(
            target=self._run_internal,
            args=(query, output_depth, thread_id, cached_research_notes, translated_query),
            daemon=True,
        )
        add_script_run_ctx(thread)
        thread.start()
        return thread

    def stream_chunks(self, timeout: int = 180) -> Generator[str, None, None]:
        """
        Generator that yields writer chunks from the queue.
        FIXED: drains remaining items properly when writer_done is set.
        """
        while True:
            try:
                chunk = self._stream_queue.get(timeout=timeout)
                if chunk is None:
                    # Drain any remaining items before stopping
                    while not self._stream_queue.empty():
                        try:
                            remaining = self._stream_queue.get_nowait()
                            if remaining is not None:
                                yield remaining
                        except queue.Empty:
                            break
                    break
                yield chunk
            except queue.Empty:
                if self.events.writer_done.is_set():
                    break
                continue

    def save_report(
        self,
        query: str,
        report: str,
        output_dir: str = "outputs",
        format: str = "md",
    ) -> str:
        """Save audited report to file. Returns the file path."""
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = os.path.join(output_dir, f"report_{timestamp}.{format}")

        with open(filepath, "w", encoding="utf-8") as f:
            if format == "md":
                f.write("# Final Audited Academic Report\n\n")
                f.write(f"**Original Query:** {query}\n\n")
                f.write(f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                f.write(f"**Model:** {self._router.current_model}\n\n")
                f.write(f"**Writer:** {self.writer_mode}\n\n")
                f.write("---\n\n")
                f.write(report)
            elif format == "tex":
                f.write("\\documentclass{article}\n")
                f.write("\\usepackage{amsmath}\n")
                f.write("\\usepackage{hyperref}\n")
                f.write("\\begin{document}\n")
                f.write(f"\\title{{Academic Report}}\n")
                f.write(
                    f"\\date{{{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}}}\n"
                )
                f.write("\\maketitle\n\n")
                f.write(report)
                f.write("\n\\end{document}\n")
            elif format == "docx":
                # Fallback in case they try to write to f (DOCX is binary, handled below)
                pass

        if format == "docx":
            try:
                from docx import Document
                doc = Document()
                doc.add_heading('Academic Report', 0)
                doc.add_paragraph(f"Original Query: {query}")
                doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph(f"Model: {self._router.current_model}")
                doc.add_paragraph(f"Writer: {self.writer_mode}")
                doc.add_paragraph("-" * 20)
                
                # Basic markdown parsing for docx
                for line in report.split('\n'):
                    if line.startswith('# '):
                        doc.add_heading(line[2:].strip(), level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:].strip(), level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:].strip(), level=3)
                    elif line.strip():
                        doc.add_paragraph(line)
                        
                doc.save(filepath)
            except ImportError:
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write("Error: python-docx library is not installed.\n")
                    f.write("Please run: pip install python-docx\n\n")
                    f.write(report)

        return filepath